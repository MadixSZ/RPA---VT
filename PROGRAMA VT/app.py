from flask import Flask, render_template, request, send_file, send_from_directory
from datetime import datetime, date, timedelta
import csv
import os
import zipfile
from io import BytesIO
import calendar

# bibliotecas opcionais (conversão para PDF / Word)
try:
    from weasyprint import HTML
except Exception:
    HTML = None

try:
    import pdfkit
except Exception:
    pdfkit = None

try:
    from docx import Document
except Exception:
    Document = None

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'recibos_html'

MESES_PT = {
    1: 'Janeiro', 2: 'Fevereiro', 3: 'Março', 4: 'Abril',
    5: 'Maio', 6: 'Junho', 7: 'Julho', 8: 'Agosto',
    9: 'Setembro', 10: 'Outubro', 11: 'Novembro', 12: 'Dezembro'
}


def converter_feriados(feriados_str, ano_referencia):
    feriados = []
    for data_str in (feriados_str or '').split(','):
        data_str = data_str.strip()
        if not data_str:
            continue
        try:
            dia, mes = map(int, data_str.split('/'))
            feriados.append(date(ano_referencia, mes, dia))
        except Exception:
            pass
    return feriados


def previous_business_day(d, feriados):
    while True:
        if d.weekday() < 5 and d not in feriados:
            return d
        d = d - timedelta(days=1)


def calcular_data_emissao(data_admissao, mes_referencia, feriados_str):
    ano, mes = map(int, mes_referencia.split('-'))
    feriados = converter_feriados(feriados_str, ano)

    primeiro_dia_mes = date(ano, mes, 1)

    if data_admissao > primeiro_dia_mes:
        data_emissao = data_admissao
    else:
        dia = 1
        while True:
            data_teste = date(ano, mes, dia)
            condicoes = [
                data_teste.weekday() < 5,
                data_teste not in feriados
            ]
            if all(condicoes):
                data_emissao = data_teste
                break
            dia += 1

    return data_emissao


def calcular_dias_uteis(data_admissao, mes_referencia, feriados_str, considerar_sabados, data_saida=None):
    ano, mes = map(int, mes_referencia.split('-'))
    primeiro_dia_mes = date(ano, mes, 1)
    data_inicio = max(data_admissao, primeiro_dia_mes)
    ultimo_dia_mes = date(ano, mes, calendar.monthrange(ano, mes)[1])

    # limita ao dia de saída se houver
    if data_saida and data_saida <= ultimo_dia_mes:
        ultimo_dia = min(ultimo_dia_mes, data_saida)
    else:
        ultimo_dia = ultimo_dia_mes

    feriados = converter_feriados(feriados_str, ano)

    # construir lista de sábados válidos a partir da data de admissão
    sabados_validos = []
    for dia in range(data_inicio.day, ultimo_dia.day + 1):
        data = date(ano, mes, dia)
        if data.weekday() == 5:
            sabados_validos.append(data)

    # aplicar alternância a partir do primeiro sábado depois da admissão
    sabados_alternados = [s for idx, s in enumerate(sabados_validos) if idx % 2 == 0] if considerar_sabados else []

    dias_uteis = 0
    for dia in range(data_inicio.day, ultimo_dia.day + 1):
        data = date(ano, mes, dia)
        condicoes = [
            data >= data_inicio,
            data not in feriados,
            (data.weekday() < 5) or (data in sabados_alternados)
        ]
        if all(condicoes):
            dias_uteis += 1

    return dias_uteis


def valor_por_extenso(valor):
    unidades = ["", "um", "dois", "três", "quatro", "cinco", 
               "seis", "sete", "oito", "nove"]
    dez_a_vinte = ["dez", "onze", "doze", "treze", "quatorze", 
                  "quinze", "dezesseis", "dezessete", "dezoito", "dezenove"]
    dezenas = ["", "dez", "vinte", "trinta", "quarenta", "cinquenta", 
              "sessenta", "setenta", "oitenta", "noventa"]
    centenas = ["", "cento", "duzentos", "trezentos", "quatrocentos", 
               "quinhentos", "seiscentos", "setecentos", "oitocentos", "novecentos"]
    
    reais = int(valor)
    centavos = int(round((valor - reais) * 100))
    
    extenso = []
    if reais >= 100:
        extenso.append(centenas[reais // 100])
        reais %= 100
    if reais >= 20:
        extenso.append(dezenas[reais // 10])
        reais %= 10
    if reais >= 10 and reais < 20:
        extenso.append(dez_a_vinte[reais - 10])
    elif reais > 0 and reais < 10:
        extenso.append(unidades[reais])

    texto = " e ".join([p for p in extenso if p])
    texto = texto if texto else "zero"
    texto += " reais"
    
    if centavos > 0:
        texto += " e "
        if centavos >= 10 and centavos < 20:
            texto += dez_a_vinte[centavos - 10]
        else:
            if centavos >= 20:
                texto += dezenas[centavos // 10]
                c = centavos % 10
                if c > 0:
                    texto += f" e {unidades[c]}"
            else:
                texto += unidades[centavos]
        texto += " centavos"
    
    return texto


@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        try:
            arquivo_csv = request.files['planilha']
            mes_referencia = request.form['mes_referencia']
            feriados = request.form.get('feriados', '')
            considerar_sabados = request.form.get('sabados_alternados') == 'on'

            funcionarios = []
            conteudo = arquivo_csv.stream.read().decode('utf-8')
            leitor = csv.DictReader(conteudo.splitlines())

            for linha in leitor:
                nome = (linha.get('nome') or '').strip()
                data_admissao_str = (linha.get('data_admissao') or '').strip()
                valor_str = (linha.get('valor_conducao') or '0').strip()

                if not data_admissao_str or not valor_str:
                    continue

                data_admissao = datetime.strptime(data_admissao_str, '%Y-%m-%d').date()
                valor = float(valor_str)

                data_saida = None
                if 'data_saida' in linha and (linha.get('data_saida') or '').strip():
                    try:
                        data_saida = datetime.strptime(linha.get('data_saida').strip(), '%Y-%m-%d').date()
                    except Exception:
                        data_saida = None

                dias = calcular_dias_uteis(data_admissao, mes_referencia, feriados, considerar_sabados, data_saida)
                data_emissao = calcular_data_emissao(data_admissao, mes_referencia, feriados)

                ano_ref, mes_ref = map(int, mes_referencia.split('-'))
                feriados_do_mes = converter_feriados(feriados, ano_ref)
                if data_saida and data_saida < data_emissao and data_saida <= date(ano_ref, mes_ref, calendar.monthrange(ano_ref, mes_ref)[1]):
                    data_emissao = previous_business_day(data_saida, feriados_do_mes)

                total_valor = dias * 2 * valor

                funcionarios.append({
                    'nome': nome,
                    'total': total_valor,
                    'total_extenso': valor_por_extenso(total_valor),
                    'dias': dias,
                    'mes_portugues': MESES_PT[mes_ref],
                    'ano_ref': ano_ref,
                    'data_emissao': data_emissao.strftime('%d de ') + MESES_PT[data_emissao.month] + data_emissao.strftime(' de %Y')
                })

            arquivos = []
            ano_pasta = funcionarios[0]['ano_ref'] if funcionarios else datetime.now().year
            mes_pasta = int(mes_referencia.split('-')[1])
            pasta_destino = os.path.join('Centro Médico', f"{mes_pasta:02d}{ano_pasta}", 'VT')
            os.makedirs(pasta_destino, exist_ok=True)

            for func in funcionarios:
                html = render_template('recibo.html', **func)

                nome_arquivo_base = f"Recibo_{func['nome'].replace(' ', '_')}"
                caminho_html = os.path.join(pasta_destino, f"{nome_arquivo_base}.html")

                with open(caminho_html, 'w', encoding='utf-8') as f:
                    f.write(html)

                arquivos.append(caminho_html)

                caminho_pdf = os.path.join(pasta_destino, f"{nome_arquivo_base}.pdf")
                try:
                    if HTML is not None:
                        HTML(string=html).write_pdf(caminho_pdf)
                    elif pdfkit is not None:
                        pdfkit.from_string(html, caminho_pdf)
                except Exception:
                    pass

                try:
                    if Document is not None:
                        doc = Document()
                        doc.add_heading('RECIBO - Entrega Vale-Transporte', level=1)
                        doc.add_paragraph(f"Empregador(a): CENTRO MÉDICO COM VIDA")
                        doc.add_paragraph(f"Empregado(a): {func['nome']}")
                        doc.add_paragraph(f"Recebi R$ {func['total']:.2f} ({func['total_extenso']})")
                        doc.add_paragraph(f"Referente a {func['dias']} dias úteis do mês de {func['mes_portugues']} de {func['ano_ref']}.")
                        doc.add_paragraph(f"São Paulo, {func['data_emissao']}.")
                        caminho_docx = os.path.join(pasta_destino, f"{nome_arquivo_base}.docx")
                        doc.save(caminho_docx)
                except Exception:
                    pass

            zip_buffer = BytesIO()
            with zipfile.ZipFile(zip_buffer, 'w') as zipf:
                for arquivo in arquivos:
                    zipf.write(arquivo, os.path.basename(arquivo))

            zip_buffer.seek(0)
            return send_file(
                zip_buffer,
                mimetype='application/zip',
                as_attachment=True,
                download_name='recibos.zip'
            )

        except Exception as e:
            return f"Erro: {str(e)}", 500

    return render_template('form.html')


@app.route('/static/<path:filename>')
def static_files(filename):
    return send_from_directory('static', filename)


if __name__ == '__main__':
    if not os.path.exists(app.config['UPLOAD_FOLDER']):
        os.makedirs(app.config['UPLOAD_FOLDER'])
    app.run(host='0.0.0.0', port=5000, debug=True)
